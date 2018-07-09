# coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

document = Document()

font_name = u'楷体_gb2312'

paragraph = document.add_paragraph()
run = paragraph.add_run(u'中文内容')
run.font.size = Pt(24)
run.font.name = font_name
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
