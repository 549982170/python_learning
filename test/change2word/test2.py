#!/usr/bin/python
# coding: UTF-8

from HTMLParser import HTMLParser
import urllib
from docx import Document
import os
import docx
import re
from docx.shared import Inches


class HTMLClient:
    #获取html网页源码
    def GetPage(self, url):
        #user_agent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
        user_agent = 'Mozilla/5.0 (X11; Linux i686) AppleWebKit/537.36 (KHTML, like Gecko) Ubuntu Chromium/34.0.1847.116 Chrome/34.0.1847.116 Safari/537.36'
        headers = { 'User-Agent' : user_agent }
        req = urllib.request.Request(url, None, headers)
        try:
            res = urllib.request.urlopen(req)
            return res.read().decode("utf-8")
        except urllib.error.HTTPError as e:
            return None
    #获取网络图片并保存在程序运行目录下
    def GetPic(self, url):
        user_agent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
        headers = { 'User-Agent' : user_agent }
        req = urllib.request.Request(url, None, headers)
        try:
            res = urllib.request.urlopen(req)
            return res.read()
        except urllib.error.HTTPError as e:
            return None

class MYHTMLParser(HTMLParser):
    def __init__(self, docfile):
        HTMLParser.__init__(self)
        self.docfile = docfile
        self.doc = Document(docfile)
        self.myclient = HTMLClient()
        self.text = ''
        self.title = False
        self.isdescription = False
        self.picList=[]
    #根据标签头类型决定标签内容的格式
    def handle_starttag(self, tag, attrs):
        #print "Encountered the beginning of a %s tag" % tag
        self.title = False
        self.isdescription = False
        #<h1>标签说明其中的内容是标题
        if re.match(r'h(\d)', tag):
            self.title = True
        #图片的处理比较复杂，首先需要找到对应的图片的url，然后下载并写入doc中
        #下载的图片格式如果有问题，docx模块会报错，因此重新定义异常处理
        #图片名称需要记录下来，在文档保存后要自动删除
        if tag == "img":
            if len(attrs) == 0: pass
            else:
                for (variable, value)  in attrs:
                    if variable == "src":
                        #此处图片url类型为[http://url/pic.img!200*200]
                        #不同网站图片类型不同，因此当作不同处理
                        picdata = self.myclient.GetPic(value.split('!')[0])
                        if picdata == None:
                            pass
                        else:
                            pictmp = value.split('/')[-1].split('!')[0]
                            picfix = value.split('/')[-1].split('!')[-1]
                            with open(pictmp, 'wb') as pic:
                                pic.write(bytes(picdata))
                                pic.close()
                            try:
                                if picfix[0:1] == 'c':
                                    self.doc.add_picture(pictmp, width=Inches(4.5))
                                else:
                                    self.doc.add_picture(pictmp)#, width=Inches(2.25))
                            except docx.image.exceptions.UnexpectedEndOfFileError as e:
                                print(e)
                            self.picList.append(pictmp)
        #javascript脚本
        if tag == 'script':
            self.isdescription = True
    def handle_data(self, data):
        if self.title == True:
            if self.text != '':
                self.doc.add_paragraph(self.text)
            self.text = ''
            self.doc.add_heading(data, level=2)
        if self.isdescription == False:
            self.text += data
    def handle_endtag(self, tag):
        #if tag == 'br' or tag == 'p' or tag == 'div':
        if self.text != '':
            self.doc.add_paragraph(self.text)
            self.text = ''
    def complete(self, html):
        self.feed(html)
        self.doc.save(self.docfile)
        for item in self.picList:
            if os.path.exists(item):
                os.remove(item)


