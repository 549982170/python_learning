# coding:utf-8
# !/user/bin/python
import os
from docx import Document
def readDocx(docName, separated_rows=0, addStr=''):
    """获取docx的文档中的所有文字,不管格式,暂不支持doc格式每line行
    @param separated_rows: int 相隔行数
    @param addStr: str 追加的字符
    """
    TotalPage = 1  # 页数
    extension = os.path.splitext(docName)[1]  # 文件后缀名
    separated_rows = int(separated_rows) # 需要隔separated_rows行追加字符
    if extension == ".txt":
        with open(docName) as f:
            if separated_rows != 0:  # 需要隔行追加字符
                docText = '<br>'.join(paragraph + addStr if index % separated_rows == 0 else paragraph for index, paragraph in enumerate(f))
                docText = docText.replace(addStr, "", 1).replace('\n', '')  # 第一行不需要分页符
                docText = docText.split("-----------------------")[0]  # 去掉页码控制符,下标为1时是页码数
                TotalPage = docText.count(addStr) + 1  # 页数统计
            else:
                docText = '<br>'.join(paragraph for paragraph in f).replace('\n', '')  # 生成迭代器,然后加入回车换行,比起生成列表需要更少内存,此情况迭代器的括号可省略
                docText = docText.split("-----------------------")[0]  # 去掉页码控制符,下标为1时是页码数
    elif extension == ".docx":
        paras = Document(docName)
        if separated_rows != 0:  # 需要隔行追加字符
            docText = '<br>'.join(paragraph.text.encode('utf-8') + addStr if index % separated_rows == 0 else paragraph.text.encode('utf-8') for index, paragraph in enumerate(paras.paragraphs))
            docText = docText.replace(addStr, "", 1)  # 第一行不需要分页符
            TotalPage = docText.count(addStr) + 1  # 页数统计
        else:
            docText = '<br>'.join(paragraph.text.encode('utf-8') for paragraph in paras.paragraphs)  # 生成迭代器,然后加入回车换行,比起生成列表需要更少内存,此情况迭代器的括号可省略
    else:
        docText = ""
    return docText, TotalPage

print readDocx("111.docx", 40, "<head>")